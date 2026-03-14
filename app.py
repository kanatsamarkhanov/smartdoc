import streamlit as st
import streamlit.components.v1 as components
from docxtpl import DocxTemplate
from io import BytesIO
import csv
import datetime
import os
import requests
import base64
import io
import pandas as pd
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import subprocess
import time

# ----------------- СОВМЕСТИМОСТЬ ПЕРЕЗАГРУЗКИ -----------------
def safe_rerun():
    """Умная функция перезагрузки для совместимости с любыми версиями Streamlit"""
    if hasattr(st, "rerun"):
        st.rerun()
    elif hasattr(st, "experimental_rerun"):
        st.experimental_rerun()

# ----------------- PAGE & SESSION -----------------
st.set_page_config(page_title="Smart Paper Generator", page_icon="📝", layout="wide")

if "lang" not in st.session_state:
    st.session_state.lang = "kz"
if "theme" not in st.session_state:
    st.session_state.theme = "light"
if "is_registered" not in st.session_state:
    st.session_state.is_registered = False
if "ui_font" not in st.session_state:
    st.session_state.ui_font = "System Default"
if "fig_count" not in st.session_state:
    st.session_state.fig_count = 1
if "tab_count" not in st.session_state:
    st.session_state.tab_count = 1
if "eq_count" not in st.session_state:
    st.session_state.eq_count = 1

# ----------------- LOCALES -----------------
locales = {
    "ru": {
        "title": "📝 Умный генератор научных статей",
        "subtitle": "Вестник ЕНУ им. Л.Н. Гумилева · Химия / География · 2025",
        "sidebar_lang": "🌍 Язык интерфейса",
        "btn_theme_dark": "🌙 Тёмная тема",
        "btn_theme_light": "☀️ Светлая тема",
        "nav_gen": "📄 Генератор статей",
        "nav_reg": "👤 Регистрация",
        "sidebar_title": "⚙️ Настройки",
        "lbl_ui_font": "Шрифт интерфейса",
        "lbl_lang": "Основной язык статьи",
        "lbl_sec": "Секция",
        "lbl_type": "Тип статьи",
        "lbl_mrnti": "МРНТИ / IRSTI",
        "sec_meta": "1. Основные метаданные",
        "lbl_title": "Название статьи",
        "lbl_authors": "Авторы",
        "lbl_authors_help": "Например: Имя Фамилия1, Имя Фамилия2",
        "lbl_affil": "Аффилиации (Место работы/учебы)",
        "lbl_affil_help": "1 Университет, Город, Страна; email",
        "lbl_email": "Email для корреспонденции",
        "sec_text": "2. Текст статьи (Загрузка IMRAD)",
        "lbl_abstract": "Аннотация (до 300 слов)",
        "lbl_kw": "Ключевые слова",
        "lbl_kw_help": "Слово 1; слово 2; слово 3 (от 3 до 10 слов)",
        "lbl_intro": "Введение (.txt/.docx)",
        "lbl_methods": "Материалы и методы (.txt/.docx)",
        "lbl_results": "Результаты (.txt/.docx)",
        "lbl_discussion": "Обсуждение (.txt/.docx)",
        "lbl_conclusion": "Заключение (.txt/.docx)",
        "lbl_ref_manager": "📚 Менеджер литературы",
        "lbl_ref_style": "Стиль цитирования",
        "lbl_fig_manager": "📊 Менеджер рисунков",
        "lbl_tab_manager": "📋 Менеджер таблиц",
        "lbl_eq_manager": "➗ Менеджер формул",
        "lbl_add_fig": "➕ Добавить рисунок",
        "lbl_add_tab": "➕ Добавить таблицу",
        "lbl_add_eq": "➕ Добавить формулу",
        "btn_upload_short": "📎 Загрузить",
        "lbl_fig_hint_title": "💡 Подсказка для графиков",
        "lbl_fig_hint_text": "Если рисунок состоит из нескольких частей (a, b, c), используйте **один тег** `[@fig1]` для всей группы.",
        "lbl_tab_hint_title": "💡 Инструкция для сложных таблиц",
        "lbl_tab_hint_text": "Для таблиц с объединенными ячейками загружайте их в формате **.docx**, чтобы сохранить форматирование.",
        "lbl_eq_hint_title": "💡 Подсказка для формул",
        "lbl_eq_hint_text": "Введите формулу. Разместите тег `[@eq1]` в тексте статьи.",
        "btn_sample_table": "📥 Скачать образец сложной таблицы",
        "lbl_samples": "📥 Скачать шаблоны файлов",
        "sec_backmatter": "4. Дополнительная информация (Back Matter)",
        "lbl_supp": "6. Supplementary Materials",
        "lbl_contrib": "7. Author Contributions",
        "lbl_auth_info": "8. Author Information",
        "lbl_funding": "9. Funding",
        "lbl_ack": "10. Acknowledgements",
        "lbl_coi": "11. Conflicts of Interest",
        "sec_trans": "3. Переводы метаданных",
        "trans_info": "По требованиям журнала необходимо предоставить название, авторов, аннотацию и ключевые слова на двух других языках.",
        "gen_btn": "🚀 Сгенерировать статью",
        "err_abs_len": "⚠️ Аннотация слишком длинная: {count} слов. Максимум: 300.",
        "succ_abs_len": "Слов в аннотации: {count}/300",
        "err_fill_req": "Пожалуйста, заполните хотя бы Название и Авторов.",
        "err_gen": "Произошла ошибка при генерации: ",
        "succ_gen": "✅ Документ успешно сгенерирован за {time} сек!",
        "btn_dl_docx": "⬇️ Скачать .docx",
        "btn_dl_pdf": "⬇️ Скачать .pdf",
        "err_pdf": "⚠️ Не удалось сгенерировать PDF (требуется LibreOffice на сервере). Доступен DOCX файл.",
        "reg_header": "📝 Регистрация исследователя",
        "reg_name": "ФИО (Полностью)",
        "reg_email": "Ваш Email",
        "reg_phone": "Номер телефона (без кода)",
        "reg_code": "Код",
        "reg_org": "Организация / Университет",
        "reg_pos": "Должность / Статус (например: Докторант)",
        "reg_submit": "Зарегистрироваться",
        "reg_success": "✅ Вы успешно зарегистрированы! Теперь вам доступен генератор статей.",
        "reg_info": "Вы можете перейти в раздел «Генератор статей».",
        "reg_req_msg": "🔒 Для создания статьи необходимо заполнить форму регистрации.",
        "reg_err_fill": "Пожалуйста, корректно заполните Имя, Email и Телефон.",
        "f_author": "Канат Самарханов / Kanat Samarkhanov",
        "f_license": "Лицензия",
        "f_univ": "ЕНУ им. Л.Н. Гумилева — Кафедра физической и экономической географии",
        "browse_files": "Выберите файл",
        "drag_drop": "Перетащите файл сюда\nПоддерживаемые форматы: txt, docx",
        "limit": "Лимит 200MB",
        "fig_prefix": "Рисунок",
        "tab_prefix": "Таблица",
        "fb_header": "💬 Оставить отзыв",
        "fb_text": "Ваши предложения или найденные ошибки",
        "fb_btn": "Отправить отзыв",
        "fb_succ": "Спасибо за ваш отзыв!"
    },
    "kz": {
        "title": "📝 Ғылыми мақалалардың ақылды генераторы",
        "subtitle": "Л.Н. Гумилев атындағы ЕҰУ Хабаршысы · Химия / География · 2025",
        "sidebar_lang": "🌍 Интерфейс тілі",
        "btn_theme_dark": "🌙 Түнгі режим",
        "btn_theme_light": "☀️ Күндізгі режим",
        "nav_gen": "📄 Мақала генераторы",
        "nav_reg": "👤 Тіркелу",
        "sidebar_title": "⚙️ Баптаулар",
        "lbl_ui_font": "Интерфейс қаріпі",
        "lbl_lang": "Мақаланың негізгі тілі",
        "lbl_sec": "Секция",
        "lbl_type": "Мақала түрі",
        "lbl_mrnti": "МРНТИ / IRSTI",
        "sec_meta": "1. Негізгі метадеректер",
        "lbl_title": "Мақаланың атауы",
        "lbl_authors": "Авторлар",
        "lbl_authors_help": "Мысалы: Аты Жөні1, Аты Жөні2",
        "lbl_affil": "Аффилиация (Жұмыс/оқу орны)",
        "lbl_affil_help": "1 Университет, Қала, Ел; email",
        "lbl_email": "Корреспонденцияға арналған email",
        "sec_text": "2. Мақала мәтіні (IMRAD Файлдары)",
        "lbl_abstract": "Аңдатпа (300 сөзге дейін)",
        "lbl_kw": "Түйінді сөздер",
        "lbl_kw_help": "Сөз 1; сөз 2; сөз 3 (3-тен 10 сөзге дейін)",
        "lbl_intro": "Кіріспе (.txt/.docx)",
        "lbl_methods": "Материалдар/әдістер (.txt/.docx)",
        "lbl_results": "Нәтижелер (.txt/.docx)",
        "lbl_discussion": "Талқылау (.txt/.docx)",
        "lbl_conclusion": "Қорытынды (.txt/.docx)",
        "lbl_ref_manager": "📚 Әдебиеттер менеджері",
        "lbl_ref_style": "Дәйексөз стилі",
        "lbl_fig_manager": "📊 Суреттер менеджері",
        "lbl_tab_manager": "📋 Кестелер менеджері",
        "lbl_eq_manager": "➗ Формулалар менеджері",
        "lbl_add_fig": "➕ Сурет қосу",
        "lbl_add_tab": "➕ Кесте қосу",
        "lbl_add_eq": "➕ Формула қосу",
        "btn_upload_short": "📎 Жүктеу",
        "lbl_fig_hint_title": "💡 Күрделі суреттер нұсқаулығы",
        "lbl_fig_hint_text": "Егер сурет бірнеше бөліктен (a, b, c) тұрса, бүкіл топ үшін **бір тегті** `[@fig1]` пайдаланыңыз.",
        "lbl_tab_hint_title": "💡 Күрделі кестелер нұсқаулығы",
        "lbl_tab_hint_text": "Кестеңіз өте кең болса немесе біріктірілген ұяшықтары болса, пішімдеуді сақтау үшін оны **.docx** форматында жүктеңіз.",
        "lbl_eq_hint_title": "💡 Формулалар нұсқаулығы",
        "lbl_eq_hint_text": "Формуланы енгізіңіз. Мәтінге `[@eq1]` тегін қойыңыз.",
        "btn_sample_table": "📥 Күрделі кесте үлгісін жүктеу",
        "lbl_samples": "📥 Файл үлгілерін жүктеп алу",
        "sec_backmatter": "4. Қосымша ақпарат (Back Matter)",
        "lbl_supp": "6. Supplementary Materials",
        "lbl_contrib": "7. Author Contributions",
        "lbl_auth_info": "8. Author Information",
        "lbl_funding": "9. Funding",
        "lbl_ack": "10. Acknowledgements",
        "lbl_coi": "11. Conflicts of Interest",
        "sec_trans": "3. Метадеректер аудармасы",
        "trans_info": "Журнал талаптарына сәйкес атауын, авторларын, аңдатпасын және түйінді сөздерін басқа екі тілде ұсыну қажет.",
        "gen_btn": "🚀 Мақаланы генерациялау",
        "err_abs_len": "⚠️ Аңдатпа тым ұзын: {count} сөз. Максимум: 300.",
        "succ_abs_len": "Аңдатпадағы сөз саны: {count}/300",
        "err_fill_req": "Кем дегенде Атауын және Авторларын толтырыңыз.",
        "err_gen": "Генерация кезінде қате пайда болды: ",
        "succ_gen": "✅ Құжат сәтті генерацияланды ({time} сек)!",
        "btn_dl_docx": "⬇️ .docx жүктеу",
        "btn_dl_pdf": "⬇️ .pdf жүктеу",
        "err_pdf": "⚠️ PDF жасау мүмкін болмады (серверде LibreOffice қажет). Тек DOCX файлы қолжетімді.",
        "reg_header": "📝 Зерттеушіні тіркеу",
        "reg_name": "Аты-жөні (Толық)",
        "reg_email": "Сіздің Email",
        "reg_phone": "Телефон нөмірі (кодсыз)",
        "reg_code": "Код",
        "reg_org": "Ұйым / Университет",
        "reg_pos": "Қызметі / Мәртебесі (мысалы: Докторант)",
        "reg_submit": "Тіркелу",
        "reg_success": "✅ Сіз жүйеге сәтті тіркелдіңіз! Енді мақала генераторы қолжетімді.",
        "reg_info": "Сіз «Мақала генераторы» бөліміне өтіп, мақала жасай аласыз.",
        "reg_req_msg": "🔒 Мақала жасау үшін тіркелу формасын толтыру қажет.",
        "reg_err_fill": "Аты-жөні, Email және Телефонды дұрыс толтырыңыз.",
        "f_author": "Канат Самарханов / Kanat Samarkhanov",
        "f_license": "Лицензия",
        "f_univ": "Л.Н. Гумилев атындағы ЕҰУ — Физикалық және экономикалық география кафедрасы",
        "browse_files": "Файлды таңдаңыз",
        "drag_drop": "Файлды осында сүйреңіз\nҚолдау көрсетілетін форматтар: txt, docx",
        "limit": "Шектеу 200MB",
        "fig_prefix": "Сурет",
        "tab_prefix": "Кесте",
        "fb_header": "💬 Кері байланыс қалдыру",
        "fb_text": "Сіздің ұсыныстарыңыз немесе табылған қателер",
        "fb_btn": "Пікір жіберу",
        "fb_succ": "Пікіріңіз үшін рақмет!"
    },
    "en": {
        "title": "📝 Smart Paper Generator",
        "subtitle": "L.N. Gumilyov ENU Bulletin · Chemistry / Geography · 2025",
        "sidebar_lang": "🌍 Language",
        "btn_theme_dark": "🌙 Dark mode",
        "btn_theme_light": "☀️ Light mode",
        "nav_gen": "📄 Paper Generator",
        "nav_reg": "👤 Registration",
        "sidebar_title": "⚙️ Paper Settings",
        "lbl_ui_font": "Interface Font",
        "lbl_lang": "Primary Language",
        "lbl_sec": "Section",
        "lbl_type": "Paper Type",
        "lbl_mrnti": "IRSTI / МРНТИ",
        "sec_meta": "1. Basic Metadata",
        "lbl_title": "Article Title",
        "lbl_authors": "Authors",
        "lbl_authors_help": "E.g.: Firstname Lastname1, Firstname Lastname2",
        "lbl_affil": "Affiliations",
        "lbl_affil_help": "1 University, City, Country; email",
        "lbl_email": "Correspondence Email",
        "sec_text": "2. Main Text (IMRAD Uploads)",
        "lbl_abstract": "Abstract (up to 300 words)",
        "lbl_kw": "Keywords",
        "lbl_kw_help": "Keyword 1; keyword 2; keyword 3 (3 to 10 words)",
        "lbl_intro": "Introduction (.txt/.docx)",
        "lbl_methods": "Materials & Methods (.txt/.docx)",
        "lbl_results": "Results (.txt/.docx)",
        "lbl_discussion": "Discussion (.txt/.docx)",
        "lbl_conclusion": "Conclusion (.txt/.docx)",
        "lbl_ref_manager": "📚 Reference Manager",
        "lbl_ref_style": "Citation Style",
        "lbl_fig_manager": "📊 Figure Manager",
        "lbl_tab_manager": "📋 Table Manager",
        "lbl_eq_manager": "➗ Equation Manager",
        "lbl_add_fig": "➕ Add Figure",
        "lbl_add_tab": "➕ Add Table",
        "lbl_add_eq": "➕ Add Equation",
        "btn_upload_short": "📎 Upload",
        "lbl_fig_hint_title": "💡 Hint for Figures",
        "lbl_fig_hint_text": "If a figure has multiple parts (a, b, c), use a **single tag** `[@fig1]`.",
        "lbl_tab_hint_title": "💡 Instruction for Complex Tables",
        "lbl_tab_hint_text": "For wide tables or tables with merged cells, please upload a **.docx** file.",
        "lbl_eq_hint_title": "💡 Equation Hint",
        "lbl_eq_hint_text": "Enter your equation. Place the tag `[@eq1]` in your text.",
        "btn_sample_table": "📥 Download Complex Table Sample",
        "lbl_samples": "📥 Download Sample Files",
        "sec_backmatter": "4. Additional Information (Back Matter)",
        "lbl_supp": "6. Supplementary Materials",
        "lbl_contrib": "7. Author Contributions",
        "lbl_auth_info": "8. Author Information",
        "lbl_funding": "9. Funding",
        "lbl_ack": "10. Acknowledgements",
        "lbl_coi": "11. Conflicts of Interest",
        "sec_trans": "3. Metadata Translations",
        "trans_info": "According to the journal requirements, the title, authors, abstract and keywords must be provided in two other languages.",
        "gen_btn": "🚀 Generate Document",
        "err_abs_len": "⚠️ Abstract is too long: {count} words. Maximum: 300.",
        "succ_abs_len": "Words in abstract: {count}/300",
        "err_fill_req": "Please fill in at least the Title and Authors.",
        "err_gen": "An error occurred during generation: ",
        "succ_gen": "✅ Document successfully generated in {time} sec!",
        "btn_dl_docx": "⬇️ Download .docx",
        "btn_dl_pdf": "⬇️ Download .pdf",
        "err_pdf": "⚠️ Failed to generate PDF (requires LibreOffice on the server). DOCX is available.",
        "reg_header": "📝 Researcher Registration",
        "reg_name": "Full Name",
        "reg_email": "Your Email",
        "reg_phone": "Phone Number (no code)",
        "reg_code": "Code",
        "reg_org": "Organization / University",
        "reg_pos": "Position / Status (e.g., PhD Student)",
        "reg_submit": "Register",
        "reg_success": "✅ You have successfully registered! The paper generator is now unlocked.",
        "reg_info": "You can now go to the 'Paper Generator' section.",
        "reg_req_msg": "🔒 To generate an article, you must complete the registration form.",
        "reg_err_fill": "Please correctly fill in your Name, Email, and Phone.",
        "f_author": "Kanat Samarkhanov",
        "f_license": "License",
        "f_univ": "L.N. Gumilyov ENU — Department of Physical and Economic Geography",
        "browse_files": "Browse files",
        "drag_drop": "Drag & drop files here\nSupported formats: txt, docx",
        "limit": "Limit 200MB",
        "fig_prefix": "Figure",
        "tab_prefix": "Table",
        "fb_header": "💬 Leave Feedback",
        "fb_text": "Your suggestions or found bugs",
        "fb_btn": "Submit Feedback",
        "fb_succ": "Thank you for your feedback!"
    }
}

l = locales[st.session_state.lang]

# ----------------- THEME CSS FROM CHECKER -----------------
dark_css = (
    "<style>"
    "html,body,[class*='css'],.stApp{background-color:#0d1b2e !important;color:#c9d8ee !important;"
    "font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Helvetica,Arial,sans-serif !important;}"
    "h1,h2,h3,h4,h5,h6,[data-testid='stMarkdownContainer'] h1,[data-testid='stMarkdownContainer'] h2,"
    "[data-testid='stMarkdownContainer'] h3{color:#e2edf7 !important;font-weight:600 !important;}"
    "p,span,label,div,li,[data-testid='stMarkdownContainer'] p,"
    "[data-testid='stCaptionContainer'],.stCaption{color:#c9d8ee !important;}"
    "[data-testid='block-container'],[data-testid='stVerticalBlock'],"
    "section[data-testid='stSidebar']{background-color:#0d1b2e !important;}"
    "[data-testid='stMetric']{background:#0f2340 !important;border:1px solid #1e3a5f !important;"
    "border-radius:6px !important;padding:12px 16px !important;}"
    "[data-testid='stMetricValue']{color:#e2edf7 !important;}"
    "[data-testid='stMetricLabel']{color:#7b96b8 !important;}"
    ".stButton>button{background-color:#0f2340 !important;color:#c9d8ee !important;"
    "border:1px solid #1e3a5f !important;border-radius:6px !important;}"
    ".stButton>button:hover{background-color:#1e3a5f !important;color:#e2edf7 !important;}"
    "[data-testid='stDownloadButton']>button{background-color:#238636 !important;color:#fff !important;"
    "border:1px solid #2ea043 !important;border-radius:6px !important;}"
    "[data-testid='stDownloadButton']>button:hover{background-color:#2ea043 !important;}"
    "[data-testid='stFileUploader']{background-color:#0f2340 !important;border-radius:8px !important;}"
    "[data-testid='stFileUploaderDropzone']{background-color:#0f2340 !important;"
    "border:2px dashed #1e3a5f !important;border-radius:8px !important;padding:24px 16px !important;}"
    "[data-testid='stFileUploaderDropzone']:hover{background-color:#112850 !important;border-color:#2f5f9e !important;}"
    "[data-testid='stFileUploader'] *,[data-testid='stFileUploaderDropzone'] *{color:#c9d8ee !important;}"
    "[data-testid='stFileUploaderDropzone'] button{background-color:#1e3a5f !important;"
    "color:#c9d8ee !important;border:1px solid #2f5f9e !important;border-radius:6px !important;"
    "padding:5px 16px !important;font-size:13px !important;font-weight:500 !important;}"
    "[data-testid='stFileUploaderDropzone'] button:hover{background-color:#2f5f9e !important;"
    "border-color:#58a6ff !important;color:#e2edf7 !important;}"
    "[data-testid='stFileUploaderFile']{background-color:#112240 !important;"
    "border:1px solid #1e3a5f !important;border-radius:6px !important;}"
    "[data-testid='stFileUploaderDeleteBtn'] button{color:#7b96b8 !important;}"
    "[data-testid='stFileUploaderDeleteBtn'] button:hover{color:#f85149 !important;}"
    "[data-testid='stDataFrame'],.stDataFrame iframe{border:1px solid #1e3a5f !important;"
    "border-radius:8px !important;overflow:hidden !important;"
    "box-shadow:0 2px 8px rgba(0,0,0,0.4) !important;}"
    "[data-testid='stAlert']{background-color:#0f2340 !important;border:1px solid #1f6feb !important;"
    "color:#c9d8ee !important;border-radius:6px !important;}"
    ".stSpinner>div{color:#c9d8ee !important;}"
    "hr{border-color:#1e3a5f !important;}"
    "input,textarea,select{background-color:#0f2340 !important;color:#c9d8ee !important;"
    "border:1px solid #1e3a5f !important;}"
    "[data-testid='stSelectbox']>div>div{background-color:#0f2340 !important;"
    "border:1px solid #1e3a5f !important;border-radius:6px !important;color:#c9d8ee !important;}"
    "</style>"
)

light_css = (
    "<style>"
    "[data-testid='stMetric']{background:#fff;padding:12px;border-radius:10px;box-shadow:0 2px 6px rgba(0,0,0,.08);}"
    "h1,h2,h3{color:#1a3a5c;}"
    "[data-testid='stDownloadButton']>button{background-color:#2ea043;color:#fff;border-radius:6px;}"
    "[data-testid='stDataFrame'],.stDataFrame iframe{border:1px solid #d0d7de;"
    "border-radius:8px;overflow:hidden;box-shadow:0 1px 4px rgba(0,0,0,0.08);}"
    "</style>"
)

st.markdown(dark_css if st.session_state.theme == "dark" else light_css, unsafe_allow_html=True)

# ----------------- HELPERS -----------------
def extract_text(uploaded_file):
    if not uploaded_file:
        return ""
    try:
        if uploaded_file.name.endswith('.txt'):
            return uploaded_file.read().decode('utf-8')
        elif uploaded_file.name.endswith('.docx'):
            doc_file = docx.Document(uploaded_file)
            return '\n'.join([p.text for p in doc_file.paragraphs])
    except Exception as e:
        return f"[Error: {str(e)}]"
    return ""

def count_wc(text):
    if not text:
        return "0 / 0"
    words = len(text.split())
    chars = len(text)
    return f"{words} / {chars}"

def create_sample_docx(section_title):
    doc = docx.Document()
    heading = doc.add_heading(section_title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Here is sample content for {section_title}. Delete this and paste your text. ")
    p.add_run("All paragraphs here are justified. ").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2 = doc.add_paragraph("Example of tagging: The results shown in [@fig1] are summarized in [@tab1]. Equation: [@eq1]. Literature supports this [@ref1].")
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_sample_table_docx():
    doc = docx.Document()
    p_tag = doc.add_paragraph("[@tab1]")
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title = doc.add_paragraph("Table 1. A complex table example with merged cells.")
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    a = table.cell(0, 0)
    b = table.cell(0, 1)
    a.merge(b)
    a.text = "Merged Header (Col 1 & 2)"
    table.cell(0, 2).text = "Header 3"
    table.cell(1, 0).text = "Data A"
    table.cell(1, 1).text = "Data B"
    table.cell(1, 2).text = "Data C"
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def append_to_github_csv(filename, row_data, header_data):
    try:
        github_token = st.secrets["GITHUB_TOKEN"]
        github_repo = st.secrets["GITHUB_REPO"]
    except Exception:
        file_exists = os.path.isfile(filename)
        with open(filename, mode="a", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            if not file_exists:
                writer.writerow(header_data)
            writer.writerow(row_data)
        return

    url = f"https://api.github.com/repos/{github_repo}/contents/{filename}"
    headers = {"Authorization": f"token {github_token}"}
    response = requests.get(url, headers=headers)
    sha = None
    if response.status_code == 200:
        data = response.json()
        sha = data["sha"]
        content = base64.b64decode(data["content"]).decode("utf-8")
    else:
        content = "\ufeff"

    output = io.StringIO()
    writer = csv.writer(output)
    if content == "\ufeff":
        writer.writerow(header_data)
    writer.writerow(row_data)
    new_content = content + output.getvalue()
    payload = {
        "message": f"Added: {filename}",
        "content": base64.b64encode(new_content.encode("utf-8")).decode("utf-8"),
    }
    if sha:
        payload["sha"] = sha
    requests.put(url, headers=headers, json=payload)

def log_generation(title_text, authors_text, lang, process_time, file_size_kb,
                   figs, tabs, refs, eqs, wc_intro, wc_meth, wc_res, wc_disc, wc_conc):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    row = [timestamp, lang, title_text, authors_text, process_time, file_size_kb,
           figs, tabs, refs, eqs, wc_intro, wc_meth, wc_res, wc_disc, wc_conc]
    header = ["Timestamp", "Language", "Title", "Authors", "Process_Time_sec",
              "File_Size_KB", "Num_Figs", "Num_Tabs", "Num_Refs", "Num_Eqs",
              "Intro(W/C)", "Methods(W/C)", "Results(W/C)", "Discussion(W/C)", "Conclusion(W/C)"]
    append_to_github_csv("generation_logs.csv", row, header)

def log_registration(name, email, phone, org, pos):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    row = [timestamp, name, email, phone, org, pos]
    header = ["Уақыты (Timestamp)", "Аты-жөні (Full Name)", "Email",
              "Телефон (Phone)", "Ұйым (Organization)", "Лауазымы (Position)"]
    append_to_github_csv("registered_users.csv", row, header)

def send_email_notification(user_email, feedback_text):
    try:
        target_email = st.secrets.get("CONTACT_EMAIL", "kanat.baurzhanuly@gmail.com")
        formsubmit_url = f"https://formsubmit.co/ajax/{target_email}"
        payload = {
            "name": "Smart Paper Generator - Жаңа пікір",
            "email": user_email if user_email else "No email provided",
            "message": feedback_text,
            "_subject": "Жаңа пікір: Smart Paper Generator (Кері байланыс)"
        }
        requests.post(formsubmit_url, json=payload)
    except Exception:
        pass

def log_feedback(user_email, feedback_text):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    row = [timestamp, user_email, feedback_text]
    header = ["Уақыты (Timestamp)", "Email", "Отзыв (Feedback)"]
    append_to_github_csv("user_feedback.csv", row, header)
    send_email_notification(user_email, feedback_text)

def convert_to_pdf(docx_path, pdf_path):
    try:
        subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', docx_path,
             '--outdir', os.path.dirname(pdf_path)],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
        if os.path.exists(pdf_path):
            return True
    except Exception:
        pass
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return True
    except Exception:
        pass
    return False

# ----------------- SIDEBAR: LANGUAGE & THEME -----------------
with st.sidebar:
    st.markdown(f"### {l['sidebar_lang']}")
    col_f1, col_f2, col_f3 = st.columns(3)

    if col_f1.button("🇰🇿", use_container_width=True):
        st.session_state.lang = "kz"
        safe_rerun()

    if col_f2.button("🇷🇺", use_container_width=True):
        st.session_state.lang = "ru"
        safe_rerun()

    if col_f3.button("🇬🇧", use_container_width=True):
        st.session_state.lang = "en"
        safe_rerun()

    st.markdown("---")

    _tbtn = l["btn_theme_light"] if st.session_state.theme == "dark" else l["btn_theme_dark"]
    if st.button(_tbtn, use_container_width=True):
        st.session_state.theme = "light" if st.session_state.theme == "dark" else "dark"
        safe_rerun()
    st.markdown("---")

# ----------------- HEADER -----------------
st.title(l["title"])
st.caption(l["subtitle"])
st.markdown("---")

if "nav_radio" not in st.session_state or st.session_state.nav_radio not in [l["nav_gen"], l["nav_reg"]]:
    st.session_state.nav_radio = l["nav_gen"]
if st.session_state.get("go_to_gen"):
    st.session_state.nav_radio = l["nav_gen"]
    st.session_state.go_to_gen = False

app_mode = st.radio("", [l["nav_gen"], l["nav_reg"]],
                    horizontal=True, label_visibility="collapsed",
                    key="nav_radio")
st.markdown("---")

is_locked = not st.session_state.is_registered

# ----------------- GENERATOR MODE -----------------
if app_mode == l["nav_gen"]:
    if is_locked:
        st.error(l["reg_req_msg"], icon="🔒")

    st.subheader(l["sidebar_title"])
    col_s1, col_s2, col_s3, col_s4, col_s5 = st.columns(5)
    with col_s1:
        primary_lang = st.selectbox(l["lbl_lang"],
                                    ["Русский", "Қазақша", "English"],
                                    disabled=is_locked)
    with col_s2:
        section = st.selectbox(l["lbl_sec"], ["Химия", "География"],
                               disabled=is_locked)
    with col_s3:
        paper_type = st.selectbox(
            l["lbl_type"],
            ["Научная статья (Article)", "Обзор (Review)",
             "Мини-обзор (Mini-review)", "Краткое сообщение (Communication)"],
            disabled=is_locked
        )
    with col_s4:
        mrnti = st.text_input(l["lbl_mrnti"], value="06.81.23",
                              disabled=is_locked)
    with col_s5:
        font_mapping = {
            "System Default": "sans-serif",
            "Times New Roman": "'Times New Roman', Times, serif",
            "Arial": "Arial, Helvetica, sans-serif",
            "Georgia": "Georgia, serif"
        }
        new_font = st.selectbox(
            l["lbl_ui_font"],
            list(font_mapping.keys()),
            index=list(font_mapping.keys()).index(st.session_state.ui_font)
        )
        if new_font != st.session_state.ui_font:
            st.session_state.ui_font = new_font
            safe_rerun()

    st.markdown("<br>", unsafe_allow_html=True)

    st.header(l["sec_meta"])
    col1, col2 = st.columns(2)
    with col1:
        title = st.text_area(l["lbl_title"], height=68, disabled=is_locked)
        authors = st.text_area(l["lbl_authors"],
                               help=l["lbl_authors_help"],
                               height=68, disabled=is_locked)
    with col2:
        affiliations = st.text_area(l["lbl_affil"],
                                    help=l["lbl_affil_help"],
                                    height=68, disabled=is_locked)
        corr_email = st.text_input(l["lbl_email"], disabled=is_locked)

    st.header(l["sec_text"])
    abstract = st.text_area(l["lbl_abstract"], height=150, disabled=is_locked)
    abstract_word_count = len(abstract.split()) if abstract else 0
    if not is_locked:
        if abstract_word_count > 300:
            st.error(l["err_abs_len"].format(count=abstract_word_count))
        elif abstract_word_count > 0:
            st.success(l["succ_abs_len"].format(count=abstract_word_count))

    keywords = st.text_input(l["lbl_kw"], help=l["lbl_kw_help"],
                             disabled=is_locked)

    st.markdown("##### " + l["lbl_samples"])
    col_dl1, col_dl2, col_dl3, col_dl4, col_dl5 = st.columns(5)
    with col_dl1:
        st.download_button("📥 Intro",
                           create_sample_docx("Introduction"),
                           file_name="sample_intro.docx",
                           use_container_width=True, disabled=is_locked)
    with col_dl2:
        st.download_button("📥 Methods",
                           create_sample_docx("Materials and Methods"),
                           file_name="sample_methods.docx",
                           use_container_width=True, disabled=is_locked)
    with col_dl3:
        st.download_button("📥 Results",
                           create_sample_docx("Results"),
                           file_name="sample_results.docx",
                           use_container_width=True, disabled=is_locked)
    with col_dl4:
        st.download_button("📥 Discussion",
                           create_sample_docx("Discussion"),
                           file_name="sample_discussion.docx",
                           use_container_width=True, disabled=is_locked)
    with col_dl5:
        st.download_button("📥 Conclusion",
                           create_sample_docx("Conclusion"),
                           file_name="sample_conclusion.docx",
                           use_container_width=True, disabled=is_locked)
    st.markdown("<br>", unsafe_allow_html=True)

    col_i1, col_i2, col_i3 = st.columns([1, 1, 1])
    with col_i1:
        file_intro = st.file_uploader(l["lbl_intro"],
                                      type=["txt", "docx"],
                                      disabled=is_locked)
        file_methods = st.file_uploader(l["lbl_methods"],
                                        type=["txt", "docx"],
                                        disabled=is_locked)
    with col_i2:
        file_results = st.file_uploader(l["lbl_results"],
                                        type=["txt", "docx"],
                                        disabled=is_locked)
        file_discussion = st.file_uploader(l["lbl_discussion"],
                                           type=["txt", "docx"],
                                           disabled=is_locked)
    with col_i3:
        file_conclusion = st.file_uploader(l["lbl_conclusion"],
                                           type=["txt", "docx"],
                                           disabled=is_locked)

    st.markdown("<br><hr>", unsafe_allow_html=True)

    # FIGURES & TABLES
    col_ft1, col_ft2 = st.columns(2)

    with col_ft1:
        st.header(l["lbl_fig_manager"])
        with st.expander(l["lbl_fig_hint_title"]):
            st.markdown(l["lbl_fig_hint_text"])

        hf1, hf2, hf3 = st.columns([1.5, 3.5, 3])
        hf1.markdown("**Tag**")
        hf2.markdown("**Caption**")
        hf3.markdown("**File**")

        for i in range(st.session_state.fig_count):
            cf1, cf2, cf3 = st.columns([1.5, 3.5, 3])
            with cf1:
                st.text_input(f"fig_tag_{i}", value=f"[@fig{i+1}]",
                              key=f"f_tag_{i}",
                              label_visibility="collapsed",
                              disabled=is_locked)
            with cf2:
                st.text_input(f"fig_cap_{i}", placeholder="Caption...",
                              key=f"f_cap_{i}",
                              label_visibility="collapsed",
                              disabled=is_locked)
            with cf3:
                st.markdown('<div class="compact-uploader"></div>',
                            unsafe_allow_html=True)
                st.file_uploader(f"fig_file_{i}",
                                 type=["png", "jpg", "jpeg"],
                                 key=f"f_file_{i}",
                                 label_visibility="collapsed",
                                 disabled=is_locked)

        if st.button(l["lbl_add_fig"], disabled=is_locked):
            st.session_state.fig_count += 1
            safe_rerun()

    with col_ft2:
        st.header(l["lbl_tab_manager"])
        with st.expander(l["lbl_tab_hint_title"]):
            st.markdown(l["lbl_tab_hint_text"])
            st.download_button(l["btn_sample_table"],
                               create_sample_table_docx(),
                               file_name="sample_complex_table.docx",
                               use_container_width=True, disabled=is_locked)

        ht1, ht2, ht3 = st.columns([1.5, 3.5, 3])
        ht1.markdown("**Tag**")
        ht2.markdown("**Caption**")
        ht3.markdown("**File**")

        for i in range(st.session_state.tab_count):
            ct1, ct2, ct3 = st.columns([1.5, 3.5, 3])
            with ct1:
                st.text_input(f"tab_tag_{i}", value=f"[@tab{i+1}]",
                              key=f"t_tag_{i}",
                              label_visibility="collapsed",
                              disabled=is_locked)
            with ct2:
                st.text_input(f"tab_cap_{i}", placeholder="Caption...",
                              key=f"t_cap_{i}",
                              label_visibility="collapsed",
                              disabled=is_locked)
            with ct3:
                st.markdown('<div class="compact-uploader"></div>',
                            unsafe_allow_html=True)
                st.file_uploader(f"tab_file_{i}",
                                 type=["xlsx", "csv", "docx", "txt"],
                                 key=f"t_file_{i}",
                                 label_visibility="collapsed",
                                 disabled=is_locked)

        if st.button(l["lbl_add_tab"], disabled=is_locked):
            st.session_state.tab_count += 1
            safe_rerun()

    st.markdown("<br>", unsafe_allow_html=True)

    # EQUATIONS
    st.header(l["lbl_eq_manager"])
    with st.expander(l["lbl_eq_hint_title"]):
        st.markdown(l["lbl_eq_hint_text"])

    he1, he2 = st.columns([1.5, 8.5])
    he1.markdown("**Tag**")
    he2.markdown("**Equation / Formula**")

    for i in range(st.session_state.eq_count):
        ce1, ce2 = st.columns([1.5, 8.5])
        with ce1:
            st.text_input(f"eq_tag_{i}", value=f"[@eq{i+1}]",
                          key=f"e_tag_{i}", label_visibility="collapsed",
                          disabled=is_locked)
        with ce2:
            st.text_input(f"eq_val_{i}", placeholder="E = mc^2 ...",
                          key=f"e_val_{i}",
                          label_visibility="collapsed",
                          disabled=is_locked)

    if st.button(l["lbl_add_eq"], disabled=is_locked):
        st.session_state.eq_count += 1
        safe_rerun()

    # REFERENCES
    st.markdown("<hr>", unsafe_allow_html=True)
    st.header(l["lbl_ref_manager"])
    ref_style = st.selectbox(l["lbl_ref_style"], ["GOST", "APA", "IEEE"],
                             disabled=is_locked)
    ref_df = pd.DataFrame([{
        "Tag in text": "[@ref1]",
        "Author(s)": "",
        "Year": "",
        "Title": "",
        "Journal/Publisher": "",
        "Volume/Pages": ""
    }])
    if not is_locked:
        edited_refs = st.data_editor(ref_df, num_rows="dynamic",
                                     use_container_width=True)
    else:
        edited_refs = ref_df
        st.dataframe(ref_df, use_container_width=True)

    # BACK MATTER
    st.header(l["sec_backmatter"])
    val_supp = st.text_area(l["lbl_supp"], value="No supplementary material.",
                            height=68, disabled=is_locked)
    val_contrib = st.text_area(
        l["lbl_contrib"],
        value=("Conceptualization, X.X. and Y.Y.; methodology, X.X.; software, X.X.; "
               "validation, X.X., Y.Y. and Z.Z.; formal analysis, X.X.; investigation, X.X.; "
               "resources, X.X.; data curation, X.X.; writing—original draft preparation, X.X.; "
               "writing—review and editing, X.X.; visualisation, X.X.; supervision, X.X.; "
               "project administration, X.X.; funding acquisition, Y.Y. All authors have read "
               "and agreed to the published version of the manuscript."),
        height=120, disabled=is_locked
    )
    val_auth_info = st.text_area(
        l["lbl_auth_info"],
        value=("Beisembayev, Adil Sayatuly - researcher, L.N. Gumilyov Eurasian "
               "National University, Kazhymukan st., 13, Astana, Kazakhstan, 010000; "
               "email: beisembayev_as@enu.kz, https://orcid.org/0001-0003-2203-9099"),
        height=80, disabled=is_locked
    )
    val_funding = st.text_area(l["lbl_funding"],
                               value="This research received no external funding.",
                               height=68, disabled=is_locked)
    val_ack = st.text_area(l["lbl_ack"],
                           value="Administrative and technical support was provided by...",
                           height=68, disabled=is_locked)
    val_coi = st.text_area(
        l["lbl_coi"],
        value=("The authors declare no conflicts of interest. The funders had no role "
               "in the study’s design, data collection, analysis, manuscript writing, "
               "or publication decisions."),
        height=80, disabled=is_locked
    )

    # TRANSLATIONS
    st.header(l["sec_trans"])
    st.info(l["trans_info"])
    trans_langs = ["Русский", "Қазақша", "English"]
    if primary_lang in trans_langs:
        trans_langs.remove(primary_lang)

    col_t1, col_t2 = st.columns(2)
    with col_t1:
        st.subheader(f"{trans_langs[0]}")
        t1_title = st.text_input(f"{l['lbl_title']} ({trans_langs[0]})",
                                 disabled=is_locked)
        t1_authors = st.text_input(f"{l['lbl_authors']} ({trans_langs[0]})",
                                   disabled=is_locked)
        t1_abstract = st.text_area(
            f"{l['lbl_abstract']} ({trans_langs[0]})",
            height=100, disabled=is_locked
        )
        t1_keywords = st.text_input(f"{l['lbl_kw']} ({trans_langs[0]})",
                                    disabled=is_locked)
    with col_t2:
        st.subheader(f"{trans_langs[1]}")
        t2_title = st.text_input(f"{l['lbl_title']} ({trans_langs[1]})",
                                 disabled=is_locked)
        t2_authors = st.text_input(f"{l['lbl_authors']} ({trans_langs[1]})",
                                   disabled=is_locked)
        t2_abstract = st.text_area(
            f"{l['lbl_abstract']} ({trans_langs[1]})",
            height=100, disabled=is_locked
        )
        t2_keywords = st.text_input(f"{l['lbl_kw']} ({trans_langs[1]})",
                                    disabled=is_locked)

    st.markdown("---")
    generate_btn = st.button(l["gen_btn"], type="primary",
                             use_container_width=True,
                             disabled=is_locked)

    if generate_btn and not is_locked:
        if abstract_word_count > 300:
            st.error(l["err_abs_len"].format(count=abstract_word_count))
        elif not title or not authors:
            st.warning(l["err_fill_req"])
        else:
            with st.spinner("Генерация документов..."):
                start_time = time.time()
                try:
                    t_intro = extract_text(file_intro)
                    t_methods = extract_text(file_methods)
                    t_results = extract_text(file_results)
                    t_discussion = extract_text(file_discussion)
                    t_conclusion = extract_text(file_conclusion)

                    wc_intro = count_wc(t_intro)
                    wc_meth = count_wc(t_methods)
                    wc_res = count_wc(t_results)
                    wc_disc = count_wc(t_discussion)
                    wc_conc = count_wc(t_conclusion)

                    main_text = ""
                    if t_intro:
                        main_text += "1. INTRODUCTION\n" + t_intro + "\n\n"
                    if t_methods:
                        main_text += "2. MATERIALS AND METHODS\n" + t_methods + "\n\n"
                    if t_results:
                        main_text += "3. RESULTS\n" + t_results + "\n\n"
                    if t_discussion:
                        main_text += "4. DISCUSSION\n" + t_discussion + "\n\n"
                    if t_conclusion:
                        main_text += "5. CONCLUSION\n" + t_conclusion + "\n\n"

                    # equations
                    added_eqs = 0
                    for i in range(st.session_state.eq_count):
                        c_tag = st.session_state.get(f"e_tag_{i}", "").strip()
                        c_val = st.session_state.get(f"e_val_{i}", "").strip()
                        if c_val:
                            if c_tag:
                                main_text = main_text.replace(c_tag, c_val)
                            added_eqs += 1

                    # figures
                    fig_text = ""
                    fig_counter = 1
                    for i in range(st.session_state.fig_count):
                        c_tag = st.session_state.get(f"f_tag_{i}", "").strip()
                        c_cap = st.session_state.get(f"f_cap_{i}", "").strip()
                        if c_cap:
                            label = f"{l['fig_prefix']} {fig_counter}"
                            fig_text += f"{label}. {c_cap}\n"
                            if c_tag:
                                main_text = main_text.replace(c_tag, label)
                            fig_counter += 1

                    # tables
                    tab_text = ""
                    tab_counter = 1
                    for i in range(st.session_state.tab_count):
                        c_tag = st.session_state.get(f"t_tag_{i}", "").strip()
                        c_cap = st.session_state.get(f"t_cap_{i}", "").strip()
                        if c_cap:
                            label = f"{l['tab_prefix']} {tab_counter}"
                            tab_text += f"{label}. {c_cap}\n"
                            if c_tag:
                                main_text = main_text.replace(c_tag, label)
                            tab_counter += 1

                    if fig_text or tab_text:
                        main_text += "\n\n--- FIGURES & TABLES ---\n" + fig_text + "\n" + tab_text

                    # back matter
                    back_matter = ""
                    if val_supp:
                        back_matter += f"6. Supplementary Materials\n{val_supp}\n\n"
                    if val_contrib:
                        back_matter += f"7. Author Contributions\n{val_contrib}\n\n"
                    if val_auth_info:
                        back_matter += f"8. Author Information\n{val_auth_info}\n\n"
                    if val_funding:
                        back_matter += f"9. Funding\n{val_funding}\n\n"
                    if val_ack:
                        back_matter += f"10. Acknowledgements\n{val_ack}\n\n"
                    if val_coi:
                        back_matter += f"11. Conflicts of Interest\n{val_coi}\n\n"
                    main_text += "\n\n" + back_matter

                    # references
                    refs_compiled = []
                    ref_counter = 1
                    for _, row in edited_refs.iterrows():
                        r_tag = str(row.get("Tag in text", "")).strip()
                        r_author = str(row.get("Author(s)", "")).strip()
                        r_year = str(row.get("Year", "")).strip()
                        r_title = str(row.get("Title", "")).strip()
                        r_journal = str(row.get("Journal/Publisher", "")).strip()
                        r_vol = str(row.get("Volume/Pages", "")).strip()

                        if not r_author or r_author == "nan":
                            continue

                        if ref_style == "APA":
                            ref_entry = f"{r_author} ({r_year}). {r_title}. {r_journal}, {r_vol}."
                            first_author = r_author.split(',')[0].strip()
                            in_text = f"({first_author} et al., {r_year})"
                        elif ref_style == "IEEE":
                            ref_entry = f"[{ref_counter}] {r_author}, \"{r_title},\" {r_journal}, {r_vol}, {r_year}."
                            in_text = f"[{ref_counter}]"
                        else:
                            ref_entry = f"{ref_counter}. {r_author} {r_title} // {r_journal}. - {r_year}. - {r_vol}."
                            in_text = f"[{ref_counter}]"

                        refs_compiled.append(ref_entry)
                        if r_tag and r_tag != "nan":
                            main_text = main_text.replace(r_tag, in_text)
                        ref_counter += 1

                    final_references = "\n".join(refs_compiled)

                    template_filename = "Russian_template_2025.docx"
                    if primary_lang == "Қазақша":
                        template_filename = "Kazakh_template_2025.docx"
                    elif primary_lang == "English":
                        template_filename = "English_template_2025.docx"

                    context = {
                        "mrnti": mrnti,
                        "section": section,
                        "paper_type": paper_type,
                        "title": title,
                        "authors": authors,
                        "affiliations": affiliations,
                        "corr_email": corr_email,
                        "abstract": abstract,
                        "keywords": keywords,
                        "main_text": main_text,
                        "references": final_references,
                        "t1_title": t1_title,
                        "t1_authors": t1_authors,
                        "t1_abstract": t1_abstract,
                        "t1_keywords": t1_keywords,
                        "t2_title": t2_title,
                        "t2_authors": t2_authors,
                        "t2_abstract": t2_abstract,
                        "t2_keywords": t2_keywords,
                    }

                    doc_tpl = DocxTemplate(template_filename)
                    doc_tpl.render(context)

                    with tempfile.TemporaryDirectory() as tmpdir:
                        docx_path = os.path.join(tmpdir, "Formatted_Article.docx")
                        pdf_path = os.path.join(tmpdir, "Formatted_Article.pdf")
                        doc_tpl.save(docx_path)
                        with open(docx_path, "rb") as f:
                            docx_bytes = f.read()
                        pdf_success = convert_to_pdf(docx_path, pdf_path)
                        pdf_bytes = None
                        if pdf_success:
                            with open(pdf_path, "rb") as f:
                                pdf_bytes = f.read()

                    process_time = round(time.time() - start_time, 2)
                    file_size_kb = round(len(docx_bytes) / 1024, 2)

                    log_generation(title, authors, primary_lang,
                                   process_time, file_size_kb,
                                   fig_counter - 1, tab_counter - 1,
                                   ref_counter - 1, added_eqs,
                                   wc_intro, wc_meth, wc_res, wc_disc, wc_conc)

                    st.success(l["succ_gen"].format(time=process_time))
                    dcol1, dcol2 = st.columns(2)
                    with dcol1:
                        st.download_button(
                            label=l["btn_dl_docx"],
                            data=docx_bytes,
                            file_name="Formatted_Article.docx",
                            mime=("application/vnd.openxmlformats-"
                                  "officedocument.wordprocessingml.document"),
                            type="primary", use_container_width=True
                        )
                    with dcol2:
                        if pdf_bytes:
                            st.download_button(
                                label=l["btn_dl_pdf"],
                                data=pdf_bytes,
                                file_name="Formatted_Article.pdf",
                                mime="application/pdf",
                                type="primary", use_container_width=True
                            )
                        else:
                            st.warning(l["err_pdf"])
                except Exception as e:
                    st.error(f"{l['err_gen']} {e}")
                    st.info("💡 Ескерту: шаблон .docx файлдары папкада болуы тиіс.")

# ----------------- REGISTRATION MODE -----------------
elif app_mode == l["nav_reg"]:
    st.header(l["reg_header"])
    if st.session_state.is_registered:
        st.success(l["reg_success"])
        st.info(l["reg_info"])
    else:
        with st.form("registration_form"):
            r_name = st.text_input(l["reg_name"])
            r_email = st.text_input(l["reg_email"])
            c1, c2 = st.columns([1, 3])
            with c1:
                country_codes = [
                    "🇰🇿 +7", "🇷🇺 +7", "🇺🇿 +998", "🇰🇬 +996",
                    "🇺🇸 +1", "🇬🇧 +44", "🇨🇳 +86"
                ]
                r_phone_code = st.selectbox(l["reg_code"], country_codes)
            with c2:
                r_phone_num = st.text_input(l["reg_phone"], max_chars=12)
            r_org = st.text_input(l["reg_org"])
            r_pos = st.text_input(l["reg_pos"])
            submitted = st.form_submit_button(l["reg_submit"], type="primary")
            if submitted:
                if r_name and r_email and len(r_phone_num) >= 7:
                    full_phone = f"{r_phone_code} {r_phone_num}"
                    with st.spinner("Тіркелу жүріп жатыр..."):
                        log_registration(r_name, r_email, full_phone, r_org, r_pos)
                    st.session_state.is_registered = True
                    st.session_state.go_to_gen = True
                    st.success(l["reg_success"])
                    safe_rerun()
                else:
                    st.error(l["reg_err_fill"])

# ----------------- FEEDBACK SECTION -----------------
st.markdown("---")
st.subheader(l["fb_header"])
with st.expander(l["fb_text"], expanded=False):
    with st.form("feedback_form", clear_on_submit=True):
        fb_email = st.text_input("Email (Optional / Міндетті емес)")
        fb_text = st.text_area("Сіздің пікіріңіз / Ваш отзыв / Your feedback",
                               height=100)
        fb_submit = st.form_submit_button(l["fb_btn"])
        if fb_submit and fb_text:
            with st.spinner("..."):
                log_feedback(fb_email, fb_text)
            st.success(l["fb_succ"])

# ----------------- ADMIN DOWNLOADS IN SIDEBAR -----------------
with st.sidebar:
    if (os.path.exists("generation_logs.csv") or
            os.path.exists("registered_users.csv") or
            os.path.exists("user_feedback.csv")):
        st.markdown("<br><br>", unsafe_allow_html=True)
        st.caption("🔒 Admin panel")
        if os.path.exists("generation_logs.csv"):
            with open("generation_logs.csv", "rb") as f:
                st.download_button(
                    label="📊 Логи генерации (.csv)",
                    data=f, file_name="generation_logs.csv",
                    mime="text/csv", use_container_width=True
                )
        if os.path.exists("registered_users.csv"):
            with open("registered_users.csv", "rb") as f:
                st.download_button(
                    label="👥 База пользователей (.csv)",
                    data=f, file_name="registered_users.csv",
                    mime="text/csv", use_container_width=True
                )
        if os.path.exists("user_feedback.csv"):
            with open("user_feedback.csv", "rb") as f:
                st.download_button(
                    label="💬 Отзывы (.csv)",
                    data=f, file_name="user_feedback.csv",
                    mime="text/csv", use_container_width=True
                )

# ----------------- FOOTER -----------------
fc = "#7b96b8" if st.session_state.theme == "dark" else "#555"
flk = "#58a6ff" if st.session_state.theme == "dark" else "#0969da"
st.markdown("---")
st.markdown(
    f'<div style="text-align:center;font-size:12px;color:{fc};'
    f'padding:12px 0 20px 0;line-height:2.2;">'
    f'<b style="font-size:13px;">© 2025 {l["f_author"]}</b><br>'
    f'📧 <a href="mailto:samarkhanov_kb@enu.kz" '
    f'style="color:{flk};text-decoration:none;">samarkhanov_kb@enu.kz</a>'
    f'&nbsp;·&nbsp;'
    f'<a href="mailto:kanat.baurzhanuly@gmail.com" '
    f'style="color:{flk};text-decoration:none;">'
    f'kanat.baurzhanuly@gmail.com</a><br>'
    f'🏛️ <a href="https://fns.enu.kz/kz/page/departments/physical-and-economical-geography/faculty-members" '
    f'target="_blank" style="color:{flk};text-decoration:none;">{l["f_univ"]}</a><br>'
    f'📄 {l["f_license"]}: '
    f'<a href="https://creativecommons.org/licenses/by/4.0/" target="_blank" '
    f'style="color:{flk};text-decoration:none;">'
    f'CC BY 4.0 — Creative Commons Attribution 4.0 International</a>'
    f'</div>',
    unsafe_allow_html=True
)
